"""
File Creators

Create various file types: CSV, images, PDFs, Word docs, etc.
"""

import io
import json
from pathlib import Path
from typing import Any
from datetime import datetime


def create_csv(
    data: list[dict] | list[list],
    headers: list[str] | None = None,
    filename: str | None = None,
) -> tuple[bytes, str]:
    """
    Create a CSV file from data.

    Args:
        data: List of dicts (keys become headers) or list of lists
        headers: Optional headers (required if data is list of lists)
        filename: Optional filename (will generate one if not provided)

    Returns:
        (csv_bytes, suggested_filename)
    """
    try:
        import pandas as pd

        if isinstance(data[0], dict):
            df = pd.DataFrame(data)
        else:
            if not headers:
                headers = [f"Column{i+1}" for i in range(len(data[0]))]
            df = pd.DataFrame(data, columns=headers)

        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False)
        csv_bytes = csv_buffer.getvalue().encode("utf-8")

    except ImportError:
        # Fallback without pandas
        import csv

        csv_buffer = io.StringIO()
        writer = csv.writer(csv_buffer)

        if isinstance(data[0], dict):
            headers = list(data[0].keys())
            writer.writerow(headers)
            for row in data:
                writer.writerow([row.get(h, "") for h in headers])
        else:
            if headers:
                writer.writerow(headers)
            for row in data:
                writer.writerow(row)

        csv_bytes = csv_buffer.getvalue().encode("utf-8")

    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"data_{timestamp}.csv"
    elif not filename.endswith(".csv"):
        filename = f"{filename}.csv"

    return csv_bytes, filename


def create_image(
    width: int = 800,
    height: int = 600,
    background: str = "white",
    elements: list[dict] | None = None,
    filename: str | None = None,
    format: str = "PNG",
) -> tuple[bytes, str]:
    """
    Create an image with optional elements.

    Args:
        width: Image width
        height: Image height
        background: Background color
        elements: List of drawing elements:
            - {"type": "text", "text": "...", "x": 100, "y": 100, "color": "black", "size": 24}
            - {"type": "rectangle", "x": 0, "y": 0, "width": 100, "height": 50, "color": "blue", "fill": True}
            - {"type": "circle", "x": 100, "y": 100, "radius": 50, "color": "red", "fill": False}
            - {"type": "line", "x1": 0, "y1": 0, "x2": 100, "y2": 100, "color": "black", "width": 2}
        filename: Optional filename
        format: Image format (PNG, JPEG, etc.)

    Returns:
        (image_bytes, suggested_filename)
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        raise ImportError("Pillow is required for image creation. Install with: pip install Pillow")

    # Create image
    img = Image.new("RGB", (width, height), color=background)
    draw = ImageDraw.Draw(img)

    # Draw elements
    if elements:
        for elem in elements:
            elem_type = elem.get("type", "")
            color = elem.get("color", "black")

            if elem_type == "text":
                x = elem.get("x", 0)
                y = elem.get("y", 0)
                text = elem.get("text", "")
                size = elem.get("size", 24)

                # Try to load a font, fall back to default
                try:
                    font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", size)
                except Exception:
                    try:
                        font = ImageFont.load_default()
                    except Exception:
                        font = None

                draw.text((x, y), text, fill=color, font=font)

            elif elem_type == "rectangle":
                x = elem.get("x", 0)
                y = elem.get("y", 0)
                w = elem.get("width", 100)
                h = elem.get("height", 100)
                fill = elem.get("fill", False)

                if fill:
                    draw.rectangle([x, y, x + w, y + h], fill=color)
                else:
                    draw.rectangle([x, y, x + w, y + h], outline=color)

            elif elem_type == "circle":
                x = elem.get("x", 0)
                y = elem.get("y", 0)
                r = elem.get("radius", 50)
                fill = elem.get("fill", False)

                if fill:
                    draw.ellipse([x - r, y - r, x + r, y + r], fill=color)
                else:
                    draw.ellipse([x - r, y - r, x + r, y + r], outline=color)

            elif elem_type == "line":
                x1 = elem.get("x1", 0)
                y1 = elem.get("y1", 0)
                x2 = elem.get("x2", 100)
                y2 = elem.get("y2", 100)
                width = elem.get("width", 1)

                draw.line([x1, y1, x2, y2], fill=color, width=width)

    # Save to bytes
    img_buffer = io.BytesIO()
    img.save(img_buffer, format=format)
    img_bytes = img_buffer.getvalue()

    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        ext = format.lower()
        filename = f"image_{timestamp}.{ext}"
    elif not filename.lower().endswith(f".{format.lower()}"):
        filename = f"{filename}.{format.lower()}"

    return img_bytes, filename


def combine_images(
    image_paths: list[str],
    layout: str = "horizontal",
    spacing: int = 10,
    background: str = "white",
    filename: str | None = None,
) -> tuple[bytes, str]:
    """
    Combine multiple images into one.

    Args:
        image_paths: List of paths to images
        layout: "horizontal" or "vertical" or "grid"
        spacing: Pixels between images
        background: Background color
        filename: Optional filename

    Returns:
        (image_bytes, suggested_filename)
    """
    try:
        from PIL import Image
    except ImportError:
        raise ImportError("Pillow is required for image operations. Install with: pip install Pillow")

    # Load all images
    images = []
    for path in image_paths:
        img = Image.open(path)
        images.append(img)

    if not images:
        raise ValueError("No images to combine")

    if layout == "horizontal":
        # Calculate total size
        total_width = sum(img.width for img in images) + spacing * (len(images) - 1)
        max_height = max(img.height for img in images)

        # Create combined image
        combined = Image.new("RGB", (total_width, max_height), color=background)

        x_offset = 0
        for img in images:
            # Center vertically
            y_offset = (max_height - img.height) // 2
            combined.paste(img, (x_offset, y_offset))
            x_offset += img.width + spacing

    elif layout == "vertical":
        max_width = max(img.width for img in images)
        total_height = sum(img.height for img in images) + spacing * (len(images) - 1)

        combined = Image.new("RGB", (max_width, total_height), color=background)

        y_offset = 0
        for img in images:
            # Center horizontally
            x_offset = (max_width - img.width) // 2
            combined.paste(img, (x_offset, y_offset))
            y_offset += img.height + spacing

    else:  # grid
        # Calculate grid dimensions (roughly square)
        n = len(images)
        cols = int(n ** 0.5)
        if cols * cols < n:
            cols += 1
        rows = (n + cols - 1) // cols

        # Get max dimensions
        max_w = max(img.width for img in images)
        max_h = max(img.height for img in images)

        total_width = cols * max_w + spacing * (cols - 1)
        total_height = rows * max_h + spacing * (rows - 1)

        combined = Image.new("RGB", (total_width, total_height), color=background)

        for i, img in enumerate(images):
            row = i // cols
            col = i % cols
            x = col * (max_w + spacing) + (max_w - img.width) // 2
            y = row * (max_h + spacing) + (max_h - img.height) // 2
            combined.paste(img, (x, y))

    # Save to bytes
    img_buffer = io.BytesIO()
    combined.save(img_buffer, format="PNG")
    img_bytes = img_buffer.getvalue()

    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"combined_{timestamp}.png"
    elif not filename.lower().endswith(".png"):
        filename = f"{filename}.png"

    return img_bytes, filename


def create_word_doc(
    content: str | list[dict],
    title: str | None = None,
    filename: str | None = None,
) -> tuple[bytes, str]:
    """
    Create a Word document.

    Args:
        content: String of text or list of elements:
            - {"type": "heading", "level": 1, "text": "..."}
            - {"type": "paragraph", "text": "..."}
            - {"type": "bullet", "items": ["...", "..."]}
            - {"type": "table", "headers": [...], "rows": [[...], [...]]}
        title: Optional document title
        filename: Optional filename

    Returns:
        (docx_bytes, suggested_filename)
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        raise ImportError("python-docx is required for Word document creation. Install with: pip install python-docx")

    doc = Document()

    if title:
        doc.add_heading(title, 0)

    if isinstance(content, str):
        # Simple text content
        for para in content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        # Structured content
        for elem in content:
            elem_type = elem.get("type", "paragraph")

            if elem_type == "heading":
                level = elem.get("level", 1)
                text = elem.get("text", "")
                doc.add_heading(text, level)

            elif elem_type == "paragraph":
                text = elem.get("text", "")
                doc.add_paragraph(text)

            elif elem_type == "bullet":
                items = elem.get("items", [])
                for item in items:
                    doc.add_paragraph(item, style="List Bullet")

            elif elem_type == "table":
                headers = elem.get("headers", [])
                rows = elem.get("rows", [])

                if headers:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = str(header)

                    for row_data in rows:
                        row_cells = table.add_row().cells
                        for i, cell_data in enumerate(row_data):
                            if i < len(row_cells):
                                row_cells[i].text = str(cell_data)

    # Save to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_bytes = doc_buffer.getvalue()

    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"document_{timestamp}.docx"
    elif not filename.lower().endswith(".docx"):
        filename = f"{filename}.docx"

    return doc_bytes, filename


def create_pdf(
    content: str | list[dict],
    title: str | None = None,
    filename: str | None = None,
) -> tuple[bytes, str]:
    """
    Create a PDF document.

    Args:
        content: String of text or list of elements similar to create_word_doc
        title: Optional document title
        filename: Optional filename

    Returns:
        (pdf_bytes, suggested_filename)
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
    except ImportError:
        raise ImportError("reportlab is required for PDF creation. Install with: pip install reportlab")

    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    if title:
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 12))

    if isinstance(content, str):
        # Simple text content
        for para in content.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), styles["Normal"]))
                story.append(Spacer(1, 6))
    else:
        # Structured content
        for elem in content:
            elem_type = elem.get("type", "paragraph")

            if elem_type == "heading":
                level = elem.get("level", 1)
                text = elem.get("text", "")
                style_name = f"Heading{min(level, 6)}"
                if style_name in styles:
                    story.append(Paragraph(text, styles[style_name]))
                else:
                    story.append(Paragraph(text, styles["Heading1"]))
                story.append(Spacer(1, 6))

            elif elem_type == "paragraph":
                text = elem.get("text", "")
                story.append(Paragraph(text, styles["Normal"]))
                story.append(Spacer(1, 6))

            elif elem_type == "bullet":
                items = elem.get("items", [])
                for item in items:
                    story.append(Paragraph(f"• {item}", styles["Normal"]))
                story.append(Spacer(1, 6))

            elif elem_type == "table":
                headers = elem.get("headers", [])
                rows = elem.get("rows", [])

                table_data = []
                if headers:
                    table_data.append(headers)
                table_data.extend(rows)

                if table_data:
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                        ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 12))

    doc.build(story)
    pdf_bytes = pdf_buffer.getvalue()

    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"document_{timestamp}.pdf"
    elif not filename.lower().endswith(".pdf"):
        filename = f"{filename}.pdf"

    return pdf_bytes, filename
